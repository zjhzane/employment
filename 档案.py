from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import os
import re

# ========== 文件路径 ==========
excel_path = r'C:\Users\zjhza\Desktop\new.xlsx'
template_path = r'C:\Users\zjhza\Desktop\高等学校毕业生档案转递单2023.docx'
output_dir = r'C:\Users\zjhza\Desktop\output_docs'
os.makedirs(output_dir, exist_ok=True)

# ========== 正则替换函数 ==========
def safe_replace(text, key, value):
    pattern = r'[ \u3000\t]*' + re.escape(key) + r'[ \u3000\t]*'
    return re.sub(pattern, value.strip(), text)

# ========== 当前日期 ==========
today = datetime.today()
month_str = str(today.month)
day_str = str(today.day)

# ========== 打开 Excel 并读取字段 ==========
wb = load_workbook(excel_path)
ws = wb.active
headers = [cell.value for cell in ws[1]]

def get_cell_value(row, col_name):
    try:
        idx = headers.index(col_name)
        return str(row[idx].value).strip() if row[idx].value is not None else ''
    except:
        return ''

# ========== 处理每一行数据 ==========
for i, row in enumerate(ws.iter_rows(min_row=2)):
    doc = Document(template_path)

    # 字段提取
    name = get_cell_value(row, '姓名')
    raw_origin_place = get_cell_value(row, '生源地名称')
    archive_type = get_cell_value(row, '档案转寄类型名称')
    id_number = get_cell_value(row, '身份证号')
    phone = get_cell_value(row, '手机号码')
    raw_unit = get_cell_value(row, '用人单位名称')
    archive_receiver = get_cell_value(row, '档案转寄单位')
    code_number = get_cell_value(row, '转递编号')

    # 用人单位处理：括号内中文
    if '（' in raw_unit and '）' in raw_unit:
        unit_name_cleaned = raw_unit.split('（')[-1].split('）')[0]
    else:
        unit_name_cleaned = raw_unit

    # 生源地处理：保留到最后一个"市"
    last_city_index = raw_origin_place.rfind('市')
    origin_cleaned = raw_origin_place[:last_city_index + 1] if last_city_index != -1 else raw_origin_place

    # 替换字段映射
    replacements = {
        '{1}': name,
        '{2}': origin_cleaned,
        '{3}': archive_type,
        '{321102200204040021}': id_number,
        '{5}': phone,
        '{6}': unit_name_cleaned,
        '{7}': archive_receiver,
        '{m}': month_str,
        '{d}': day_str,
        '{2510876CYLAXKAZFVR}': code_number,
    }

    # ========== 替换段落内容 ==========
    for paragraph in doc.paragraphs:
        text = paragraph.text
        replaced = text
        for key, val in replacements.items():
            replaced = safe_replace(replaced, key, val)
        if replaced != text:
            for run in paragraph.runs:
                run.text = ''
            new_run = paragraph.add_run(replaced)
            if paragraph.runs:
                ref = paragraph.runs[0]
                new_run.font.name = ref.font.name
                new_run.font.size = ref.font.size
                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), ref.font.name)

    # ========== 替换表格内容 ==========
    for table in doc.tables:
        for row_cells in table.rows:
            for cell in row_cells.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    replaced = text
                    for key, val in replacements.items():
                        replaced = safe_replace(replaced, key, val)
                    if replaced != text:
                        for run in paragraph.runs:
                            run.text = ''
                        new_run = paragraph.add_run(replaced)
                        if paragraph.runs:
                            ref = paragraph.runs[0]
                            new_run.font.name = ref.font.name
                            new_run.font.size = ref.font.size
                            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), ref.font.name)

    # ========== 去除"生源地"下一格首个空格（run 拼接情况下也处理） ==========
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                if cell.text.strip() == '生源地':
                    if row_index + 1 < len(table.rows):
                        target_cell = table.cell(row_index + 1, col_index)
                        for para in target_cell.paragraphs:
                            if not para.runs:
                                continue

                            # 保存第一个run的格式
                            ref_run = para.runs[0]
                            font_name = ref_run.font.name
                            font_size = ref_run.font.size

                            # 获取并处理文本
                            full_text = ''.join(run.text for run in para.runs)
                            stripped_text = full_text.lstrip()

                            # 如果文本有变化，则清空段落并重新写入
                            if full_text != stripped_text:
                                para.clear()
                                new_run = para.add_run(stripped_text)
                                # 恢复字体格式
                                new_run.font.name = font_name
                                new_run.font.size = font_size
                                if font_name:
                                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # ========== 保存输出文件 ==========
    output_path = os.path.join(output_dir, f'档案转递单_{name or f"unknown_{i+1}"}.docx')
    doc.save(output_path)
    print(f'✅ 已生成：{output_path}')
